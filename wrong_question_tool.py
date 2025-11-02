#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
错题整理工具
作者：mmm
版本：2.0.0
支持平台：Windows, Linux, macOS, Android (通过 Termux)
赞助链接：https://gitee.com/orangearc655743/Wrong-Question-Tool.git
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import os
import shutil
import json
import threading
import datetime
from PIL import Image, ImageTk, ImageEnhance, ImageFilter, ImageDraw
import pytesseract
import cv2
import numpy as np
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches
import sys
import platform
import webbrowser

class WrongQuestionTool:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("错题整理工具 v2.0.0 - 作者：mmm")
        self.root.geometry("1400x900")
        
        # 设置主题样式
        self.setup_style()
        
        # 预设学科
        self.subjects = ["语文", "数学", "英语", "物理", "化学", "政治", "历史", "生物", "地理", "其它"]
        
        # 程序目录
        self.program_dir = os.path.dirname(os.path.abspath(__file__))
        self.cuoti_dir = os.path.join(self.program_dir, "CuoTi")
        self.config_file = os.path.join(self.program_dir, "config.json")
        self.themes_dir = os.path.join(self.program_dir, "themes")
        
        # 创建必要目录
        for dir_path in [self.cuoti_dir, self.themes_dir]:
            if not os.path.exists(dir_path):
                os.makedirs(dir_path)
        
        # 加载配置
        self.config = self.load_config()
        
        # 当前路径
        self.current_path = self.cuoti_dir
        self.path_history = [self.cuoti_dir]
        
        # 搜索相关
        self.search_var = tk.StringVar()
        self.search_results = []
        
        # 统计信息
        self.stats = {"total_files": 0, "total_size": 0, "by_subject": {}}
        
        # 快捷键绑定
        self.setup_shortcuts()
        
        self.setup_ui()
        self.refresh_file_list()
        self.update_stats()
        
    def setup_style(self):
        """设置主题样式"""
        self.style = ttk.Style()
        
        # 检查可用的主题
        available_themes = self.style.theme_names()
        if 'clam' in available_themes:
            self.style.theme_use('clam')
        elif 'alt' in available_themes:
            self.style.theme_use('alt')
        
        # 配置样式
        self.style.configure('Title.TLabel', font=('微软雅黑', 16, 'bold'), foreground='#2c3e50')
        self.style.configure('Header.TLabel', font=('微软雅黑', 12, 'bold'), foreground='#34495e')
        self.style.configure('Custom.TButton', font=('微软雅黑', 10))
        self.style.configure('Custom.Treeview', font=('微软雅黑', 9))
        self.style.configure('Custom.Treeview.Heading', font=('微软雅黑', 10, 'bold'))
    
    def setup_shortcuts(self):
        """设置快捷键"""
        self.root.bind('<Control-i>', lambda e: self.import_questions())
        self.root.bind('<Control-r>', lambda e: self.refresh_file_list())
        self.root.bind('<Control-f>', lambda e: self.focus_search())
        self.root.bind('<Control-e>', lambda e: self.export_pdf())
        self.root.bind('<F5>', lambda e: self.refresh_file_list())
        self.root.bind('<Delete>', lambda e: self.delete_item())
        self.root.bind('<F2>', lambda e: self.rename_item())
    
    def load_config(self):
        """加载配置文件"""
        default_config = {
            "last_subject": "语文",
            "ocr_enabled": True,
            "image_quality": 90,
            "export_format": "pdf",
            "theme": "default",
            "auto_backup": True,
            "show_stats": True,
            "image_rotation": 0,
            "crop_settings": {"left": 0, "top": 0, "right": 100, "bottom": 100}
        }
        
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    # 合并默认配置
                    for key, value in default_config.items():
                        if key not in config:
                            config[key] = value
                    return config
        except Exception as e:
            print(f"加载配置失败: {e}")
        
        return default_config
    
    def save_config(self):
        """保存配置文件"""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存配置失败: {e}")
    
    def setup_ui(self):
        """设置用户界面"""
        # 创建主菜单
        self.create_menu()
        
        # 创建工具栏
        self.create_toolbar()
        
        # 创建主内容区域
        self.create_main_content()
        
        # 创建状态栏
        self.create_status_bar()
    
    def create_menu(self):
        """创建菜单栏"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # 文件菜单
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="文件", menu=file_menu)
        file_menu.add_command(label="导入错题", accelerator="Ctrl+I", command=self.import_questions)
        file_menu.add_command(label="刷新列表", accelerator="Ctrl+R", command=self.refresh_file_list)
        file_menu.add_separator()
        file_menu.add_command(label="导出为PDF", accelerator="Ctrl+E", command=self.export_pdf)
        file_menu.add_command(label="导出为Word", command=self.export_word)
        file_menu.add_command(label="导出选中项", command=self.export_selected)
        file_menu.add_separator()
        file_menu.add_command(label="批量重命名", command=self.batch_rename)
        file_menu.add_command(label="创建备份", command=self.create_backup)
        file_menu.add_separator()
        file_menu.add_command(label="退出", command=self.root.quit)
        
        # 编辑菜单
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="编辑", menu=edit_menu)
        edit_menu.add_command(label="重命名", accelerator="F2", command=self.rename_item)
        edit_menu.add_command(label="删除", accelerator="Delete", command=self.delete_item)
        edit_menu.add_command(label="复制", command=self.copy_item)
        edit_menu.add_command(label="移动", command=self.move_item)
        edit_menu.add_separator()
        edit_menu.add_command(label="添加标签", command=self.add_tags)
        edit_menu.add_command(label="编辑标签", command=self.edit_tags)
        
        # 工具菜单
        tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="工具", menu=tools_menu)
        tools_menu.add_command(label="图片预处理", command=self.image_preprocessing)
        tools_menu.add_command(label="图片裁剪", command=self.image_cropping)
        tools_menu.add_command(label="图片旋转", command=self.image_rotation)
        tools_menu.add_command(label="OCR识别", command=self.ocr_recognition)
        tools_menu.add_separator()
        tools_menu.add_command(label="批量处理", command=self.batch_process)
        tools_menu.add_command(label="搜索文件", accelerator="Ctrl+F", command=self.show_search)
        
        # 视图菜单
        view_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="视图", menu=view_menu)
        view_menu.add_command(label="显示统计信息", command=self.show_statistics)
        view_menu.add_command(label="切换主题", command=self.switch_theme)
        view_menu.add_separator()
        view_menu.add_command(label="大图标", command=lambda: self.change_view_mode("large"))
        view_menu.add_command(label="小图标", command=lambda: self.change_view_mode("small"))
        view_menu.add_command(label="列表", command=lambda: self.change_view_mode("list"))
        
        # 帮助菜单
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="帮助", menu=help_menu)
        help_menu.add_command(label="使用说明", command=self.show_help)
        help_menu.add_command(label="快捷键", command=self.show_shortcuts)
        help_menu.add_command(label="关于", command=self.show_about)
        help_menu.add_command(label="赞助链接", command=self.open_sponsor_link)
    
    def create_toolbar(self):
        """创建工具栏"""
        toolbar = ttk.Frame(self.root)
        toolbar.pack(side=tk.TOP, fill=tk.X, padx=5, pady=2)
        
        # 左侧按钮组
        left_frame = ttk.Frame(toolbar)
        left_frame.pack(side=tk.LEFT)
        
        # 导入按钮
        ttk.Button(left_frame, text="导入错题", command=self.import_questions, style='Custom.TButton').pack(side=tk.LEFT, padx=2)
        
        # 刷新按钮
        ttk.Button(left_frame, text="刷新", command=self.refresh_file_list, style='Custom.TButton').pack(side=tk.LEFT, padx=2)
        
        # 返回上一级按钮
        ttk.Button(left_frame, text="返回上一级", command=self.go_up, style='Custom.TButton').pack(side=tk.LEFT, padx=2)
        
        # 搜索框
        search_frame = ttk.Frame(toolbar)
        search_frame.pack(side=tk.LEFT, padx=20)
        
        ttk.Label(search_frame, text="搜索:").pack(side=tk.LEFT)
        search_entry = ttk.Entry(search_frame, textvariable=self.search_var, width=20)
        search_entry.pack(side=tk.LEFT, padx=5)
        search_entry.bind('<Return>', lambda e: self.perform_search())
        search_entry.bind('<KeyRelease>', lambda e: self.live_search())
        
        ttk.Button(search_frame, text="搜索", command=self.perform_search, style='Custom.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(search_frame, text="清除", command=self.clear_search, style='Custom.TButton').pack(side=tk.LEFT, padx=2)
        
        # 右侧按钮组
        right_frame = ttk.Frame(toolbar)
        right_frame.pack(side=tk.RIGHT)
        
        # 导出按钮
        ttk.Button(right_frame, text="导出PDF", command=self.export_pdf, style='Custom.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(right_frame, text="导出Word", command=self.export_word, style='Custom.TButton').pack(side=tk.LEFT, padx=2)
        
        # 工具按钮
        tools_menu = ttk.Menubutton(right_frame, text="工具")
        tools_menu.pack(side=tk.LEFT, padx=2)
        tools_dropdown = tk.Menu(tools_menu, tearoff=0)
        tools_dropdown.add_command(label="图片预处理", command=self.image_preprocessing)
        tools_dropdown.add_command(label="图片裁剪", command=self.image_cropping)
        tools_dropdown.add_command(label="图片旋转", command=self.image_rotation)
        tools_dropdown.add_command(label="OCR识别", command=self.ocr_recognition)
        tools_menu.config(menu=tools_dropdown)
        
        # 设置按钮
        ttk.Button(right_frame, text="设置", command=self.show_settings, style='Custom.TButton').pack(side=tk.LEFT, padx=2)
    
    def create_main_content(self):
        """创建主内容区域"""
        # 创建paned window
        paned = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True, padx=5, pady=2)
        
        # 左侧：错题列表
        left_frame = ttk.Frame(paned)
        paned.add(left_frame, weight=2)
        
        # 路径和统计信息
        header_frame = ttk.Frame(left_frame)
        header_frame.pack(fill=tk.X, pady=2)
        
        # 当前路径显示
        path_frame = ttk.Frame(header_frame)
        path_frame.pack(fill=tk.X, pady=2)
        
        ttk.Label(path_frame, text="当前路径:", style='Header.TLabel').pack(side=tk.LEFT)
        self.path_var = tk.StringVar(value=self.current_path)
        path_label = ttk.Label(path_frame, textvariable=self.path_var, foreground="#3498db")
        path_label.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        # 统计信息
        stats_frame = ttk.Frame(header_frame)
        stats_frame.pack(fill=tk.X, pady=2)
        
        self.stats_var = tk.StringVar(value="统计信息加载中...")
        ttk.Label(stats_frame, textvariable=self.stats_var, style='Header.TLabel').pack(side=tk.LEFT)
        
        # 文件列表
        list_frame = ttk.Frame(left_frame)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建Treeview
        columns = ("名称", "类型", "大小", "修改时间", "标签")
        self.file_tree = ttk.Treeview(list_frame, columns=columns, show="tree headings", style='Custom.Treeview')
        
        # 设置列标题
        self.file_tree.heading("#0", text="")
        self.file_tree.heading("名称", text="名称")
        self.file_tree.heading("类型", text="类型")
        self.file_tree.heading("大小", text="大小")
        self.file_tree.heading("修改时间", text="修改时间")
        self.file_tree.heading("标签", text="标签")
        
        # 设置列宽
        self.file_tree.column("#0", width=0, stretch=tk.NO)
        self.file_tree.column("名称", width=200)
        self.file_tree.column("类型", width=80)
        self.file_tree.column("大小", width=100)
        self.file_tree.column("修改时间", width=150)
        self.file_tree.column("标签", width=100)
        
        # 添加滚动条
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.file_tree.yview)
        self.file_tree.configure(yscrollcommand=scrollbar.set)
        
        self.file_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 绑定事件
        self.file_tree.bind("<Double-1>", self.on_item_double_click)
        self.file_tree.bind("<Button-3>", self.show_context_menu)
        self.file_tree.bind("<<TreeviewSelect>>", self.on_item_select)
        
        # 右侧：预览和编辑区域
        right_frame = ttk.Frame(paned)
        paned.add(right_frame, weight=1)
        
        # 预览标签页
        notebook = ttk.Notebook(right_frame)
        notebook.pack(fill=tk.BOTH, expand=True, pady=2)
        
        # 图片预览标签页
        preview_frame = ttk.Frame(notebook)
        notebook.add(preview_frame, text="图片预览")
        
        # 图片显示区域
        self.image_frame = ttk.Frame(preview_frame)
        self.image_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.image_label = ttk.Label(self.image_frame, text="选择图片进行预览", style='Title.TLabel')
        self.image_label.pack(expand=True)
        
        # 编辑标签页
        edit_frame = ttk.Frame(notebook)
        notebook.add(edit_frame, text="编辑信息")
        
        # 编辑表单
        form_frame = ttk.Frame(edit_frame)
        form_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(form_frame, text="题目名称:", style='Header.TLabel').grid(row=0, column=0, sticky=tk.W, pady=5)
        self.title_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.title_var, width=40).grid(row=0, column=1, sticky=tk.W+tk.E, pady=5)
        
        ttk.Label(form_frame, text="学科:", style='Header.TLabel').grid(row=1, column=0, sticky=tk.W, pady=5)
        self.subject_var = tk.StringVar(value=self.config.get("last_subject", "语文"))
        subject_combo = ttk.Combobox(form_frame, textvariable=self.subject_var, values=self.subjects, width=37)
        subject_combo.grid(row=1, column=1, sticky=tk.W+tk.E, pady=5)
        
        ttk.Label(form_frame, text="标签:", style='Header.TLabel').grid(row=2, column=0, sticky=tk.W, pady=5)
        self.tags_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.tags_var, width=40).grid(row=2, column=1, sticky=tk.W+tk.E, pady=5)
        
        ttk.Label(form_frame, text="备注:", style='Header.TLabel').grid(row=3, column=0, sticky=tk.W+tk.N, pady=5)
        self.notes_text = tk.Text(form_frame, width=40, height=8)
        self.notes_text.grid(row=3, column=1, sticky=tk.W+tk.E+tk.N+tk.S, pady=5)
        
        # 按钮
        button_frame = ttk.Frame(edit_frame)
        button_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Button(button_frame, text="保存编辑", command=self.save_edit, style='Custom.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="OCR识别", command=self.ocr_recognition, style='Custom.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="图片处理", command=self.image_preprocessing, style='Custom.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="图片裁剪", command=self.image_cropping, style='Custom.TButton').pack(side=tk.LEFT, padx=5)
        
        # 配置grid权重
        form_frame.columnconfigure(1, weight=1)
        edit_frame.rowconfigure(3, weight=1)
        edit_frame.columnconfigure(0, weight=1)
    
    def create_status_bar(self):
        """创建状态栏"""
        self.status_bar = ttk.Frame(self.root)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # 左侧状态信息
        left_frame = ttk.Frame(self.status_bar)
        left_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        self.status_var = tk.StringVar(value="就绪")
        status_label = ttk.Label(left_frame, textvariable=self.status_var)
        status_label.pack(side=tk.LEFT, padx=5)
        
        # 右侧信息
        right_frame = ttk.Frame(self.status_bar)
        right_frame.pack(side=tk.RIGHT)
        
        self.selection_var = tk.StringVar(value="未选择文件")
        selection_label = ttk.Label(right_frame, textvariable=self.selection_var)
        selection_label.pack(side=tk.RIGHT, padx=5)
        
        # 进度条
        self.progress = ttk.Progressbar(self.status_bar, mode='indeterminate')
        self.progress.pack(side=tk.BOTTOM, fill=tk.X)
    
    def refresh_file_list(self):
        """刷新文件列表"""
        self.status_var.set("正在刷新...")
        self.progress.start()
        
        # 清空现有项目
        for item in self.file_tree.get_children():
            self.file_tree.delete(item)
        
        try:
            # 添加返回上一级项
            if self.current_path != self.cuoti_dir:
                parent_dir = os.path.dirname(self.current_path)
                if parent_dir.startswith(self.cuoti_dir):
                    self.file_tree.insert("", "end", text="..", values=("返回上一级", "文件夹", "", "", ""), tags=("folder", "parent"))
            
            # 添加文件和文件夹
            items = os.listdir(self.current_path)
            items.sort()
            
            for item in items:
                if item.startswith('.'):
                    continue
                    
                item_path = os.path.join(self.current_path, item)
                if os.path.isdir(item_path):
                    # 文件夹
                    self.file_tree.insert("", "end", text=item, values=(item, "文件夹", "", "", ""), tags=("folder",))
                else:
                    # 文件
                    file_size = os.path.getsize(item_path)
                    if file_size < 1024:
                        size_str = f"{file_size} B"
                    elif file_size < 1024*1024:
                        size_str = f"{file_size/1024:.1f} KB"
                    else:
                        size_str = f"{file_size/(1024*1024):.1f} MB"
                    
                    mod_time = os.path.getmtime(item_path)
                    time_str = datetime.datetime.fromtimestamp(mod_time).strftime("%Y-%m-%d %H:%M")
                    
                    file_type = self.get_file_type(item)
                    tags = self.get_file_tags(item)
                    self.file_tree.insert("", "end", text=item, values=(item, file_type, size_str, time_str, tags), tags=("file",))
            
            self.path_var.set(self.current_path)
            self.status_var.set(f"已加载 {len(items)} 个项目")
            self.update_stats()
            
        except Exception as e:
            messagebox.showerror("错误", f"刷新失败: {str(e)}")
            self.status_var.set("刷新失败")
        
        finally:
            self.progress.stop()
    
    def get_file_type(self, filename):
        """获取文件类型"""
        ext = os.path.splitext(filename)[1].lower()
        type_map = {
            '.jpg': '图片',
            '.jpeg': '图片',
            '.png': '图片',
            '.bmp': '图片',
            '.gif': '图片',
            '.tiff': '图片',
            '.pdf': 'PDF文档',
            '.doc': 'Word文档',
            '.docx': 'Word文档',
            '.txt': '文本文件',
            '.meta': '元数据'
        }
        return type_map.get(ext, '未知类型')
    
    def get_file_tags(self, filename):
        """获取文件标签"""
        try:
            meta_file = os.path.join(self.current_path, f"{os.path.splitext(filename)[0]}.meta")
            if os.path.exists(meta_file):
                with open(meta_file, 'r', encoding='utf-8') as f:
                    metadata = json.load(f)
                    return metadata.get('tags', '')
        except:
            pass
        return ''
    
    def update_stats(self):
        """更新统计信息"""
        try:
            total_files = 0
            total_size = 0
            by_subject = {}
            
            for root, dirs, files in os.walk(self.cuoti_dir):
                for file in files:
                    if not file.startswith('.'):
                        file_path = os.path.join(root, file)
                        if os.path.isfile(file_path):
                            total_files += 1
                            total_size += os.path.getsize(file_path)
                            
                            # 按学科统计
                            subject = os.path.basename(root)
                            if subject not in by_subject:
                                by_subject[subject] = 0
                            by_subject[subject] += 1
            
            self.stats = {
                "total_files": total_files,
                "total_size": total_size,
                "by_subject": by_subject
            }
            
            # 格式化大小
            if total_size < 1024:
                size_str = f"{total_size} B"
            elif total_size < 1024*1024:
                size_str = f"{total_size/1024:.1f} KB"
            elif total_size < 1024*1024*1024:
                size_str = f"{total_size/(1024*1024):.1f} MB"
            else:
                size_str = f"{total_size/(1024*1024*1024):.1f} GB"
            
            stats_text = f"总计: {total_files} 个文件, {size_str}"
            self.stats_var.set(stats_text)
            
        except Exception as e:
            self.stats_var.set("统计信息加载失败")
    
    def focus_search(self):
        """聚焦搜索框"""
        # 找到搜索框并聚焦
        for child in self.root.winfo_children():
            if isinstance(child, ttk.Frame):
                for subchild in child.winfo_children():
                    if isinstance(subchild, ttk.Entry):
                        subchild.focus_set()
                        break
    
    def perform_search(self):
        """执行搜索"""
        search_term = self.search_var.get().strip()
        if not search_term:
            self.clear_search()
            return
        
        self.status_var.set(f"搜索中: {search_term}")
        self.search_results = []
        
        # 清空现有选择
        for item in self.file_tree.get_children():
            self.file_tree.delete(item)
        
        try:
            # 递归搜索
            for root, dirs, files in os.walk(self.current_path):
                for file in files:
                    if search_term.lower() in file.lower():
                        file_path = os.path.join(root, file)
                        relative_path = os.path.relpath(file_path, self.current_path)
                        
                        file_size = os.path.getsize(file_path)
                        if file_size < 1024:
                            size_str = f"{file_size} B"
                        elif file_size < 1024*1024:
                            size_str = f"{file_size/1024:.1f} KB"
                        else:
                            size_str = f"{file_size/(1024*1024):.1f} MB"
                        
                        mod_time = os.path.getmtime(file_path)
                        time_str = datetime.datetime.fromtimestamp(mod_time).strftime("%Y-%m-%d %H:%M")
                        
                        file_type = self.get_file_type(file)
                        tags = self.get_file_tags(file)
                        
                        self.file_tree.insert("", "end", text=file, 
                                             values=(file, file_type, size_str, time_str, tags), 
                                             tags=("file",))
                        self.search_results.append(file_path)
            
            if self.search_results:
                self.status_var.set(f"找到 {len(self.search_results)} 个匹配项")
            else:
                self.status_var.set("未找到匹配项")
                
        except Exception as e:
            messagebox.showerror("搜索错误", f"搜索失败: {str(e)}")
    
    def live_search(self):
        """实时搜索"""
        # 可以在这里实现实时搜索功能
        pass
    
    def clear_search(self):
        """清除搜索"""
        self.search_var.set("")
        self.refresh_file_list()
        self.status_var.set("就绪")
    
    def on_item_double_click(self, event):
        """双击事件"""
        selection = self.file_tree.selection()
        if not selection:
            return
        
        item = selection[0]
        values = self.file_tree.item(item, "values")
        
        if values[1] == "文件夹" or values[0] == "返回上一级":
            if values[0] == "返回上一级":
                self.go_up()
            else:
                # 进入子文件夹
                folder_name = values[0]
                new_path = os.path.join(self.current_path, folder_name)
                if os.path.exists(new_path) and os.path.isdir(new_path):
                    self.path_history.append(self.current_path)
                    self.current_path = new_path
                    self.refresh_file_list()
        else:
            # 打开文件
            file_path = os.path.join(self.current_path, values[0])
            self.preview_file(file_path)
    
    def go_up(self):
        """返回上一级"""
        if len(self.path_history) > 1:
            self.path_history.pop()
            self.current_path = self.path_history[-1]
            self.refresh_file_list()
    
    def preview_file(self, file_path):
        """预览文件"""
        try:
            if file_path.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif')):
                # 图片预览
                self.show_image_preview(file_path)
            elif file_path.lower().endswith('.txt'):
                # 文本预览
                self.show_text_preview(file_path)
            else:
                self.image_label.config(text=f"不支持预览文件类型: {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("预览错误", f"无法预览文件: {str(e)}")
    
    def show_image_preview(self, image_path):
        """显示图片预览"""
        try:
            # 加载图片
            image = Image.open(image_path)
            
            # 计算显示尺寸
            max_width = 400
            max_height = 300
            
            # 计算缩放比例
            width, height = image.size
            ratio = min(max_width/width, max_height/height)
            new_width = int(width * ratio)
            new_height = int(height * ratio)
            
            # 调整大小
            image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            # 转换为tkinter格式
            photo = ImageTk.PhotoImage(image)
            
            # 显示图片
            self.image_label.config(image=photo, text="")
            self.image_label.image = photo  # 保持引用
            
            # 更新编辑信息
            self.title_var.set(os.path.splitext(os.path.basename(image_path))[0])
            
        except Exception as e:
            self.image_label.config(image="", text=f"无法加载图片: {str(e)}")
    
    def show_text_preview(self, text_path):
        """显示文本预览"""
        try:
            with open(text_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            self.image_label.config(image="", text="文本文件预览:\n\n" + content[:500] + ("..." if len(content) > 500 else ""))
            
            # 更新编辑信息
            self.title_var.set(os.path.splitext(os.path.basename(text_path))[0])
            self.notes_text.delete(1.0, tk.END)
            self.notes_text.insert(1.0, content)
            
        except Exception as e:
            self.image_label.config(image="", text=f"无法读取文本文件: {str(e)}")
    
    def import_questions(self):
        """导入错题"""
        # 选择学科
        subject = simpledialog.askstring("选择学科", "请输入学科名称:", 
                                       initialvalue=self.config.get("last_subject", "语文"),
                                       parent=self.root)
        if not subject:
            return
        
        if subject not in self.subjects:
            if messagebox.askyesno("添加学科", f"学科 '{subject}' 不在预设列表中，是否添加到列表？"):
                self.subjects.append(subject)
            else:
                return
        
        # 更新配置
        self.config["last_subject"] = subject
        self.save_config()
        
        # 选择文件
        filetypes = [
            ("图片文件", "*.jpg *.jpeg *.png *.bmp *.gif *.tiff"),
            ("所有文件", "*.*")
        ]
        
        files = filedialog.askopenfilenames(
            title="选择错题图片",
            filetypes=filetypes,
            parent=self.root
        )
        
        if not files:
            return
        
        # 创建学科目录
        subject_dir = os.path.join(self.cuoti_dir, subject)
        if not os.path.exists(subject_dir):
            os.makedirs(subject_dir)
        
        # 复制文件
        success_count = 0
        for file_path in files:
            try:
                filename = os.path.basename(file_path)
                # 确保文件名唯一
                name, ext = os.path.splitext(filename)
                counter = 1
                new_filename = filename
                while os.path.exists(os.path.join(subject_dir, new_filename)):
                    new_filename = f"{name}_{counter}{ext}"
                    counter += 1
                
                # 复制文件
                dest_path = os.path.join(subject_dir, new_filename)
                shutil.copy2(file_path, dest_path)
                success_count += 1
                
            except Exception as e:
                messagebox.showwarning("导入警告", f"导入文件 {filename} 失败: {str(e)}")
        
        if success_count > 0:
            messagebox.showinfo("导入完成", f"成功导入 {success_count} 个错题文件")
            # 刷新列表
            self.current_path = subject_dir
            self.path_history = [self.cuoti_dir, subject_dir]
            self.refresh_file_list()
        else:
            messagebox.showwarning("导入失败", "没有成功导入任何文件")
    
    def on_item_select(self, event):
        """项目选择事件"""
        selection = self.file_tree.selection()
        if selection:
            item = selection[0]
            values = self.file_tree.item(item, "values")
            
            # 更新状态栏
            if values[0] == "返回上一级":
                self.selection_var.set("返回上一级")
            elif values[1] == "文件夹":
                self.selection_var.set(f"文件夹: {values[0]}")
            else:
                self.selection_var.set(f"文件: {values[0]} ({values[2]})")
            
            # 预览文件
            if values[1] != "文件夹" and values[0] != "返回上一级":
                file_path = os.path.join(self.current_path, values[0])
                self.preview_file(file_path)
        else:
            self.selection_var.set("未选择文件")
    
    def show_context_menu(self, event):
        """显示右键菜单"""
        item = self.file_tree.identify_row(event.y)
        if item:
            self.file_tree.selection_set(item)
            context_menu = tk.Menu(self.root, tearoff=0)
            
            # 基本操作
            context_menu.add_command(label="重命名", accelerator="F2", command=self.rename_item)
            context_menu.add_command(label="删除", accelerator="Del", command=self.delete_item)
            context_menu.add_command(label="复制", command=self.copy_item)
            context_menu.add_command(label="移动", command=self.move_item)
            context_menu.add_separator()
            
            # 标签操作
            context_menu.add_command(label="添加标签", command=self.add_tags)
            context_menu.add_command(label="编辑标签", command=self.edit_tags)
            context_menu.add_separator()
            
            # 图片操作（仅对图片文件显示）
            filename = self.file_tree.item(item, "values")[0]
            if filename.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif')):
                context_menu.add_command(label="图片预处理", command=self.image_preprocessing)
                context_menu.add_command(label="图片裁剪", command=self.image_cropping)
                context_menu.add_command(label="图片旋转", command=self.image_rotation)
                context_menu.add_command(label="OCR识别", command=self.ocr_recognition)
                context_menu.add_separator()
            
            # 导出操作
            context_menu.add_command(label="导出选中项", command=self.export_selected)
            context_menu.add_command(label="导出为PDF", command=self.export_pdf)
            context_menu.add_command(label="导出为Word", command=self.export_word)
            
            try:
                context_menu.tk_popup(event.x_root, event.y_root)
            finally:
                context_menu.grab_release()
    
    def rename_item(self):
        """重命名项目"""
        selection = self.file_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一个项目")
            return
        
        item = selection[0]
        old_name = self.file_tree.item(item, "values")[0]
        
        if old_name == "返回上一级":
            return
        
        new_name = simpledialog.askstring("重命名", "请输入新名称:", initialvalue=old_name, parent=self.root)
        if not new_name or new_name == old_name:
            return
        
        old_path = os.path.join(self.current_path, old_name)
        new_path = os.path.join(self.current_path, new_name)
        
        if os.path.exists(new_path):
            messagebox.showerror("错误", "目标名称已存在")
            return
        
        try:
            os.rename(old_path, new_path)
            self.refresh_file_list()
            self.status_var.set(f"已重命名: {old_name} -> {new_name}")
        except Exception as e:
            messagebox.showerror("错误", f"重命名失败: {str(e)}")
    
    def delete_item(self):
        """删除项目"""
        selection = self.file_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一个项目")
            return
        
        item = selection[0]
        name = self.file_tree.item(item, "values")[0]
        
        if name == "返回上一级":
            return
        
        if not messagebox.askyesno("确认删除", f"确定要删除 '{name}' 吗？此操作不可恢复！"):
            return
        
        item_path = os.path.join(self.current_path, name)
        
        try:
            if os.path.isdir(item_path):
                shutil.rmtree(item_path)
            else:
                os.remove(item_path)
            
            self.refresh_file_list()
            self.status_var.set(f"已删除: {name}")
        except Exception as e:
            messagebox.showerror("错误", f"删除失败: {str(e)}")
    
    def copy_item(self):
        """复制项目"""
        selection = self.file_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一个项目")
            return
        
        item = selection[0]
        name = self.file_tree.item(item, "values")[0]
        
        if name == "返回上一级":
            return
        
        source_path = os.path.join(self.current_path, name)
        target_dir = filedialog.askdirectory(title="选择复制目标目录", parent=self.root)
        
        if not target_dir:
            return
        
        try:
            target_path = os.path.join(target_dir, name)
            # 确保目标文件名唯一
            if os.path.exists(target_path):
                name_parts = os.path.splitext(name)
                counter = 1
                while os.path.exists(target_path):
                    new_name = f"{name_parts[0]}_copy{counter}{name_parts[1]}"
                    target_path = os.path.join(target_dir, new_name)
                    counter += 1
            
            if os.path.isdir(source_path):
                shutil.copytree(source_path, target_path)
            else:
                shutil.copy2(source_path, target_path)
            
            self.status_var.set(f"已复制到: {target_path}")
            messagebox.showinfo("复制完成", f"已复制到: {target_path}")
        except Exception as e:
            messagebox.showerror("错误", f"复制失败: {str(e)}")
    
    def move_item(self):
        """移动项目"""
        selection = self.file_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一个项目")
            return
        
        item = selection[0]
        name = self.file_tree.item(item, "values")[0]
        
        if name == "返回上一级":
            return
        
        source_path = os.path.join(self.current_path, name)
        target_dir = filedialog.askdirectory(title="选择移动目标目录", parent=self.root)
        
        if not target_dir:
            return
        
        try:
            target_path = os.path.join(target_dir, name)
            # 确保目标文件名唯一
            if os.path.exists(target_path):
                name_parts = os.path.splitext(name)
                counter = 1
                while os.path.exists(target_path):
                    new_name = f"{name_parts[0]}_move{counter}{name_parts[1]}"
                    target_path = os.path.join(target_dir, new_name)
                    counter += 1
            
            shutil.move(source_path, target_path)
            self.refresh_file_list()
            self.status_var.set(f"已移动到: {target_path}")
            messagebox.showinfo("移动完成", f"已移动到: {target_path}")
        except Exception as e:
            messagebox.showerror("错误", f"移动失败: {str(e)}")
    
    def save_edit(self):
        """保存编辑"""
        selection = self.file_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一个文件")
            return
        
        item = selection[0]
        filename = self.file_tree.item(item, "values")[0]
        
        if filename == "返回上一级":
            return
        
        title = self.title_var.get().strip()
        subject = self.subject_var.get().strip()
        tags = self.tags_var.get().strip()
        notes = self.notes_text.get(1.0, tk.END).strip()
        
        if not title:
            messagebox.showwarning("警告", "请输入题目名称")
            return
        
        # 保存到元数据文件
        base_name = os.path.splitext(filename)[0]
        metadata_file = os.path.join(self.current_path, f"{base_name}.meta")
        
        metadata = {
            "title": title,
            "subject": subject,
            "tags": tags,
            "notes": notes,
            "modified_time": datetime.datetime.now().isoformat()
        }
        
        try:
            # 如果元数据文件已存在，保留原有数据
            if os.path.exists(metadata_file):
                with open(metadata_file, 'r', encoding='utf-8') as f:
                    existing_metadata = json.load(f)
                metadata.update(existing_metadata)
            
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)
            
            self.status_var.set("编辑已保存")
            messagebox.showinfo("保存成功", "编辑信息已保存")
            self.refresh_file_list()
        except Exception as e:
            messagebox.showerror("错误", f"保存失败: {str(e)}")
    
    def image_preprocessing(self):
        """图片预处理"""
        selection = self.file_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一个图片文件")
            return
        
        item = selection[0]
        filename = self.file_tree.item(item, "values")[0]
        
        if not filename.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif')):
            messagebox.showwarning("警告", "请选择一个图片文件")
            return
        
        # 创建预处理窗口
        preprocess_window = tk.Toplevel(self.root)
        preprocess_window.title("图片预处理")
        preprocess_window.geometry("600x500")
        preprocess_window.transient(self.root)
        preprocess_window.grab_set()
        
        # 图片显示区域
        img_frame = ttk.Frame(preprocess_window)
        img_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 原始图片
        original_frame = ttk.LabelFrame(img_frame, text="原始图片")
        original_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
        
        original_label = ttk.Label(original_frame)
        original_label.pack(expand=True)
        
        # 处理后图片
        processed_frame = ttk.LabelFrame(img_frame, text="处理后图片")
        processed_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5)
        
        processed_label = ttk.Label(processed_frame)
        processed_label.pack(expand=True)
        
        # 控制面板
        control_frame = ttk.Frame(preprocess_window)
        control_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # 加载图片
        image_path = os.path.join(self.current_path, filename)
        original_image = Image.open(image_path)
        
        # 显示原始图片
        self.show_image_in_label(original_image, original_label)
        
        # 处理参数
        brightness_var = tk.DoubleVar(value=1.0)
        contrast_var = tk.DoubleVar(value=1.0)
        sharpness_var = tk.DoubleVar(value=1.0)
        
        def update_preview():
            try:
                # 应用处理
                processed = original_image.copy()
                
                # 亮度调整
                if brightness_var.get() != 1.0:
                    enhancer = ImageEnhance.Brightness(processed)
                    processed = enhancer.enhance(brightness_var.get())
                
                # 对比度调整
                if contrast_var.get() != 1.0:
                    enhancer = ImageEnhance.Contrast(processed)
                    processed = enhancer.enhance(contrast_var.get())
                
                # 锐化调整
                if sharpness_var.get() != 1.0:
                    enhancer = ImageEnhance.Sharpness(processed)
                    processed = enhancer.enhance(sharpness_var.get())
                
                # 显示处理后的图片
                self.show_image_in_label(processed, processed_label)
                
            except Exception as e:
                print(f"预览更新失败: {e}")
        
        # 创建控制滑块
        ttk.Label(control_frame, text="亮度:").grid(row=0, column=0, sticky=tk.W, padx=5)
        brightness_scale = ttk.Scale(control_frame, from_=0.1, to=2.0, 
                                   variable=brightness_var, orient=tk.HORIZONTAL,
                                   command=lambda x: update_preview())
        brightness_scale.grid(row=0, column=1, sticky=tk.W+tk.E, padx=5)
        
        ttk.Label(control_frame, text="对比度:").grid(row=1, column=0, sticky=tk.W, padx=5)
        contrast_scale = ttk.Scale(control_frame, from_=0.1, to=2.0,
                                 variable=contrast_var, orient=tk.HORIZONTAL,
                                 command=lambda x: update_preview())
        contrast_scale.grid(row=1, column=1, sticky=tk.W+tk.E, padx=5)
        
        ttk.Label(control_frame, text="锐化:").grid(row=2, column=0, sticky=tk.W, padx=5)
        sharpness_scale = ttk.Scale(control_frame, from_=0.1, to=3.0,
                                  variable=sharpness_var, orient=tk.HORIZONTAL,
                                  command=lambda x: update_preview())
        sharpness_scale.grid(row=2, column=1, sticky=tk.W+tk.E, padx=5)
        
        # 按钮
        button_frame = ttk.Frame(control_frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="应用并保存", 
                  command=lambda: self.apply_image_processing(image_path, original_image, 
                                                             brightness_var.get(), contrast_var.get(), 
                                                             sharpness_var.get(), preprocess_window)).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", 
                  command=preprocess_window.destroy).pack(side=tk.LEFT, padx=5)
        
        # 配置grid权重
        control_frame.columnconfigure(1, weight=1)
        
        # 初始化预览
        update_preview()
    
    def show_image_in_label(self, image, label):
        """在标签中显示图片"""
        try:
            # 计算显示尺寸
            max_width = 250
            max_height = 200
            
            width, height = image.size
            ratio = min(max_width/width, max_height/height)
            new_width = int(width * ratio)
            new_height = int(height * ratio)
            
            image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
            photo = ImageTk.PhotoImage(image)
            
            label.config(image=photo, text="")
            label.image = photo
        except Exception as e:
            label.config(image="", text=f"无法显示图片: {str(e)}")
    
    def apply_image_processing(self, image_path, original_image, brightness, contrast, sharpness, window):
        """应用图片处理并保存"""
        try:
            # 应用处理
            processed = original_image.copy()
            
            if brightness != 1.0:
                enhancer = ImageEnhance.Brightness(processed)
                processed = enhancer.enhance(brightness)
            
            if contrast != 1.0:
                enhancer = ImageEnhance.Contrast(processed)
                processed = enhancer.enhance(contrast)
            
            if sharpness != 1.0:
                enhancer = ImageEnhance.Sharpness(processed)
                processed = enhancer.enhance(sharpness)
            
            # 备份原图（如果启用自动备份）
            if self.config.get("auto_backup", True):
                backup_path = image_path + ".backup"
                if not os.path.exists(backup_path):
                    shutil.copy2(image_path, backup_path)
            
            # 确保目录存在
            os.makedirs(os.path.dirname(image_path), exist_ok=True)
            
            # 保存处理后的图片
            processed.save(image_path, quality=self.config.get("image_quality", 90), optimize=True)
            
            messagebox.showinfo("处理完成", "图片处理完成，已保存")
            window.destroy()
            
            # 刷新预览
            self.refresh_file_list()
            
            # 如果当前预览的就是这个图片，更新预览
            selection = self.file_tree.selection()
            if selection:
                item = selection[0]
                filename = self.file_tree.item(item, "values")[0]
                if filename == os.path.basename(image_path):
                    self.preview_file(image_path)
            
        except Exception as e:
            messagebox.showerror("错误", f"图片处理失败: {str(e)}")
    
    def image_cropping(self):
        """图片裁剪功能"""
        selection = self.file_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一个图片文件")
            return
        
        item = selection[0]
        filename = self.file_tree.item(item, "values")[0]
        
        if not filename.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif')):
            messagebox.showwarning("警告", "请选择一个图片文件")
            return
        
        # 创建裁剪窗口
        crop_window = tk.Toplevel(self.root)
        crop_window.title(f"图片裁剪 - {filename}")
        crop_window.geometry("800x600")
        crop_window.transient(self.root)
        crop_window.grab_set()
        
        # 图片显示区域
        img_frame = ttk.Frame(crop_window)
        img_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 原始图片
        original_frame = ttk.LabelFrame(img_frame, text="原始图片")
        original_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
        
        original_label = ttk.Label(original_frame)
        original_label.pack(expand=True)
        
        # 裁剪后图片
        cropped_frame = ttk.LabelFrame(img_frame, text="裁剪预览")
        cropped_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5)
        
        cropped_label = ttk.Label(cropped_frame)
        cropped_label.pack(expand=True)
        
        # 裁剪控制
        control_frame = ttk.Frame(crop_window)
        control_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # 加载图片
        image_path = os.path.join(self.current_path, filename)
        original_image = Image.open(image_path)
        
        # 显示原始图片
        self.show_image_in_label(original_image, original_label, 350, 250)
        
        # 裁剪参数
        crop_vars = {
            'left': tk.IntVar(value=0),
            'top': tk.IntVar(value=0),
            'right': tk.IntVar(value=original_image.width),
            'bottom': tk.IntVar(value=original_image.height)
        }
        
        def update_crop_preview():
            try:
                left = crop_vars['left'].get()
                top = crop_vars['top'].get()
                right = crop_vars['right'].get()
                bottom = crop_vars['bottom'].get()
                
                # 确保裁剪区域有效
                if left >= right or top >= bottom:
                    return
                
                # 裁剪图片
                cropped = original_image.crop((left, top, right, bottom))
                
                # 显示裁剪后的图片
                self.show_image_in_label(cropped, cropped_label, 350, 250)
                
            except Exception as e:
                print(f"裁剪预览更新失败: {e}")
        
        # 创建裁剪控制滑块
        ttk.Label(control_frame, text="左:").grid(row=0, column=0, sticky=tk.W, padx=5)
        left_scale = ttk.Scale(control_frame, from_=0, to=original_image.width, 
                             variable=crop_vars['left'], orient=tk.HORIZONTAL,
                             command=lambda x: update_crop_preview())
        left_scale.grid(row=0, column=1, sticky=tk.W+tk.E, padx=5)
        
        ttk.Label(control_frame, text="上:").grid(row=1, column=0, sticky=tk.W, padx=5)
        top_scale = ttk.Scale(control_frame, from_=0, to=original_image.height,
                            variable=crop_vars['top'], orient=tk.HORIZONTAL,
                            command=lambda x: update_crop_preview())
        top_scale.grid(row=1, column=1, sticky=tk.W+tk.E, padx=5)
        
        ttk.Label(control_frame, text="右:").grid(row=2, column=0, sticky=tk.W, padx=5)
        right_scale = ttk.Scale(control_frame, from_=0, to=original_image.width,
                              variable=crop_vars['right'], orient=tk.HORIZONTAL,
                              command=lambda x: update_crop_preview())
        right_scale.grid(row=2, column=1, sticky=tk.W+tk.E, padx=5)
        
        ttk.Label(control_frame, text="下:").grid(row=3, column=0, sticky=tk.W, padx=5)
        bottom_scale = ttk.Scale(control_frame, from_=0, to=original_image.height,
                               variable=crop_vars['bottom'], orient=tk.HORIZONTAL,
                               command=lambda x: update_crop_preview())
        bottom_scale.grid(row=3, column=1, sticky=tk.W+tk.E, padx=5)
        
        # 预设裁剪选项
        preset_frame = ttk.LabelFrame(control_frame, text="预设裁剪")
        preset_frame.grid(row=4, column=0, columnspan=2, pady=10)
        
        def apply_preset(preset_type):
            width, height = original_image.size
            if preset_type == "square":
                # 正方形裁剪（取较小边长）
                size = min(width, height)
                left = (width - size) // 2
                top = (height - size) // 2
                right = left + size
                bottom = top + size
            elif preset_type == "top":
                # 顶部区域
                left, top = 0, 0
                right, bottom = width, height // 2
            elif preset_type == "center":
                # 中心区域
                left, top = width // 4, height // 4
                right, bottom = 3 * width // 4, 3 * height // 4
            
            crop_vars['left'].set(left)
            crop_vars['top'].set(top)
            crop_vars['right'].set(right)
            crop_vars['bottom'].set(bottom)
            update_crop_preview()
        
        ttk.Button(preset_frame, text="正方形", command=lambda: apply_preset("square")).pack(side=tk.LEFT, padx=5)
        ttk.Button(preset_frame, text="顶部", command=lambda: apply_preset("top")).pack(side=tk.LEFT, padx=5)
        ttk.Button(preset_frame, text="中心", command=lambda: apply_preset("center")).pack(side=tk.LEFT, padx=5)
        
        # 按钮
        button_frame = ttk.Frame(control_frame)
        button_frame.grid(row=5, column=0, columnspan=2, pady=10)
        
        def apply_crop():
            try:
                left = crop_vars['left'].get()
                top = crop_vars['top'].get()
                right = crop_vars['right'].get()
                bottom = crop_vars['bottom'].get()
                
                if left >= right or top >= bottom:
                    messagebox.showerror("错误", "裁剪区域无效")
                    return
                
                # 备份原图
                if self.config.get("auto_backup", True):
                    backup_path = image_path + ".backup"
                    if not os.path.exists(backup_path):
                        shutil.copy2(image_path, backup_path)
                
                # 裁剪并保存
                cropped = original_image.crop((left, top, right, bottom))
                cropped.save(image_path, quality=self.config.get("image_quality", 90), optimize=True)
                
                messagebox.showinfo("裁剪完成", "图片裁剪完成，已保存")
                crop_window.destroy()
                
                # 刷新预览
                self.refresh_file_list()
                
            except Exception as e:
                messagebox.showerror("错误", f"图片裁剪失败: {str(e)}")
        
        ttk.Button(button_frame, text="应用裁剪", command=apply_crop).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=crop_window.destroy).pack(side=tk.LEFT, padx=5)
        
        # 配置grid权重
        control_frame.columnconfigure(1, weight=1)
        
        # 初始化预览
        update_crop_preview()
    
    def image_rotation(self):
        """图片旋转功能"""
        selection = self.file_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一个图片文件")
            return
        
        item = selection[0]
        filename = self.file_tree.item(item, "values")[0]
        
        if not filename.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif')):
            messagebox.showwarning("警告", "请选择一个图片文件")
            return
        
        # 旋转选项窗口
        rotation_window = tk.Toplevel(self.root)
        rotation_window.title(f"图片旋转 - {filename}")
        rotation_window.geometry("400x300")
        rotation_window.transient(self.root)
        rotation_window.grab_set()
        
        ttk.Label(rotation_window, text="选择旋转角度:", style='Header.TLabel').pack(pady=20)
        
        # 旋转选项
        angle_var = tk.IntVar(value=90)
        
        ttk.Radiobutton(rotation_window, text="90° 顺时针", variable=angle_var, value=90).pack(anchor=tk.W, padx=20, pady=5)
        ttk.Radiobutton(rotation_window, text="180°", variable=angle_var, value=180).pack(anchor=tk.W, padx=20, pady=5)
        ttk.Radiobutton(rotation_window, text="270° 顺时针", variable=angle_var, value=270).pack(anchor=tk.W, padx=20, pady=5)
        ttk.Radiobutton(rotation_window, text="自定义角度", variable=angle_var, value=0).pack(anchor=tk.W, padx=20, pady=5)
        
        custom_frame = ttk.Frame(rotation_window)
        custom_frame.pack(pady=10)
        
        ttk.Label(custom_frame, text="角度:").pack(side=tk.LEFT)
        custom_angle = ttk.Spinbox(custom_frame, from_=0, to=360, width=10)
        custom_angle.pack(side=tk.LEFT, padx=5)
        
        def apply_rotation():
            try:
                image_path = os.path.join(self.current_path, filename)
                
                # 加载图片
                image = Image.open(image_path)
                
                # 确定旋转角度
                if angle_var.get() == 0:
                    rotation_angle = int(custom_angle.get())
                else:
                    rotation_angle = angle_var.get()
                
                # 旋转图片
                rotated = image.rotate(rotation_angle, expand=True, fillcolor='white')
                
                # 备份原图
                if self.config.get("auto_backup", True):
                    backup_path = image_path + ".backup"
                    if not os.path.exists(backup_path):
                        shutil.copy2(image_path, backup_path)
                
                # 保存旋转后的图片
                rotated.save(image_path, quality=self.config.get("image_quality", 90), optimize=True)
                
                messagebox.showinfo("旋转完成", f"图片已旋转 {rotation_angle}°")
                rotation_window.destroy()
                
                # 刷新预览
                self.refresh_file_list()
                
            except Exception as e:
                messagebox.showerror("错误", f"图片旋转失败: {str(e)}")
        
        button_frame = ttk.Frame(rotation_window)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text="应用旋转", command=apply_rotation).pack(side=tk.LEFT, padx=10)
        ttk.Button(button_frame, text="取消", command=rotation_window.destroy).pack(side=tk.LEFT, padx=10)
    
    def show_image_in_label(self, image, label, max_width=250, max_height=200):
        """在标签中显示图片（带尺寸控制）"""
        try:
            width, height = image.size
            ratio = min(max_width/width, max_height/height)
            new_width = int(width * ratio)
            new_height = int(height * ratio)
            
            image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
            photo = ImageTk.PhotoImage(image)
            
            label.config(image=photo, text="")
            label.image = photo
        except Exception as e:
            label.config(image="", text=f"无法显示图片: {str(e)}")
    
    def ocr_recognition(self):
        """OCR文字识别"""
        selection = self.file_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一个图片文件")
            return
        
        item = selection[0]
        filename = self.file_tree.item(item, "values")[0]
        
        if not filename.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif')):
            messagebox.showwarning("警告", "请选择一个图片文件")
            return
        
        self.status_var.set("正在进行OCR识别...")
        self.progress.start()
        
        def ocr_thread():
            try:
                image_path = os.path.join(self.current_path, filename)
                
                # 加载图片
                image = cv2.imread(image_path)
                if image is None:
                    raise ValueError("无法加载图片")
                
                # 预处理图片
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
                
                # 应用阈值
                _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                
                # OCR识别
                text = pytesseract.image_to_string(thresh, lang='chi_sim+eng')
                
                # 在主线程中更新UI
                self.root.after(0, lambda: self.show_ocr_result(text, filename))
                
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("OCR错误", f"OCR识别失败: {str(e)}"))
                self.root.after(0, lambda: self.status_var.set("OCR识别失败"))
            finally:
                self.root.after(0, lambda: self.progress.stop())
        
        # 在新线程中运行OCR
        thread = threading.Thread(target=ocr_thread)
        thread.daemon = True
        thread.start()
    
    def show_ocr_result(self, text, filename):
        """显示OCR结果"""
        self.status_var.set("OCR识别完成")
        
        # 创建结果窗口
        result_window = tk.Toplevel(self.root)
        result_window.title(f"OCR识别结果 - {filename}")
        result_window.geometry("600x400")
        result_window.transient(self.root)
        result_window.grab_set()
        
        # 文本显示区域
        text_frame = ttk.Frame(result_window)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        ttk.Label(text_frame, text="识别到的文字:").pack(anchor=tk.W)
        
        text_widget = tk.Text(text_frame, wrap=tk.WORD)
        text_widget.pack(fill=tk.BOTH, expand=True, pady=5)
        text_widget.insert(1.0, text)
        
        # 按钮
        button_frame = ttk.Frame(result_window)
        button_frame.pack(fill=tk.X, padx=10, pady=5)
        
        def save_text():
            text_content = text_widget.get(1.0, tk.END).strip()
            if text_content:
                # 保存为文本文件
                base_name = os.path.splitext(filename)[0]
                text_file = os.path.join(self.current_path, f"{base_name}_ocr.txt")
                
                try:
                    with open(text_file, 'w', encoding='utf-8') as f:
                        f.write(text_content)
                    
                    messagebox.showinfo("保存成功", f"OCR结果已保存到: {text_file}")
                    result_window.destroy()
                    self.refresh_file_list()
                except Exception as e:
                    messagebox.showerror("保存失败", f"保存OCR结果失败: {str(e)}")
            else:
                messagebox.showwarning("警告", "没有可保存的内容")
        
        def copy_to_clipboard():
            text_content = text_widget.get(1.0, tk.END).strip()
            if text_content:
                self.root.clipboard_clear()
                self.root.clipboard_append(text_content)
                messagebox.showinfo("复制成功", "OCR结果已复制到剪贴板")
            else:
                messagebox.showwarning("警告", "没有可复制的内容")
        
        ttk.Button(button_frame, text="保存到文件", command=save_text).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="复制到剪贴板", command=copy_to_clipboard).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="关闭", command=result_window.destroy).pack(side=tk.RIGHT, padx=5)
    
    def export_pdf(self):
        """导出PDF"""
        # 选择导出范围
        export_window = tk.Toplevel(self.root)
        export_window.title("导出PDF")
        export_window.geometry("400x300")
        export_window.transient(self.root)
        export_window.grab_set()
        
        ttk.Label(export_window, text="选择导出范围:").pack(pady=10)
        
        # 导出选项
        export_var = tk.StringVar(value="current")
        
        ttk.Radiobutton(export_window, text="当前目录所有文件", 
                       variable=export_var, value="current").pack(anchor=tk.W, padx=20, pady=5)
        ttk.Radiobutton(export_window, text="按学科导出", 
                       variable=export_var, value="subject").pack(anchor=tk.W, padx=20, pady=5)
        ttk.Radiobutton(export_window, text="选择文件和文件夹", 
                       variable=export_var, value="selected").pack(anchor=tk.W, padx=20, pady=5)
        
        # 学科选择（当选择按学科导出时显示）
        subject_frame = ttk.Frame(export_window)
        subject_frame.pack(fill=tk.X, padx=20, pady=5)
        
        ttk.Label(subject_frame, text="学科:").pack(side=tk.LEFT)
        subject_var = tk.StringVar()
        subject_combo = ttk.Combobox(subject_frame, textvariable=subject_var, 
                                   values=self.subjects, width=20)
        subject_combo.pack(side=tk.LEFT, padx=5)
        
        def on_export_option_change(*args):
            if export_var.get() == "subject":
                subject_frame.pack(fill=tk.X, padx=20, pady=5)
            else:
                subject_frame.pack_forget()
        
        export_var.trace('w', on_export_option_change)
        
        # 按钮
        button_frame = ttk.Frame(export_window)
        button_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def do_export():
            export_type = export_var.get()
            
            if export_type == "current":
                export_paths = [self.current_path]
            elif export_type == "subject":
                subject = subject_var.get()
                if not subject:
                    messagebox.showwarning("警告", "请选择学科")
                    return
                export_paths = [os.path.join(self.cuoti_dir, subject)]
            else:  # selected
                # 让用户选择文件和文件夹
                export_paths = filedialog.askopenfilenames(
                    title="选择要导出的文件",
                    initialdir=self.current_path,
                    parent=export_window
                )
                if not export_paths:
                    return
            
            # 选择保存位置
            output_file = filedialog.asksaveasfilename(
                title="保存PDF文件",
                defaultextension=".pdf",
                filetypes=[("PDF文件", "*.pdf")],
                parent=export_window
            )
            
            if not output_file:
                return
            
            export_window.destroy()
            self.perform_pdf_export(export_paths, output_file)
        
        ttk.Button(button_frame, text="确定", command=do_export).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=export_window.destroy).pack(side=tk.RIGHT, padx=5)
    
    def perform_pdf_export(self, export_paths, output_file):
        """执行PDF导出"""
        self.status_var.set("正在导出PDF...")
        self.progress.start()
        
        def export_thread():
            try:
                c = canvas.Canvas(output_file, pagesize=A4)
                width, height = A4
                
                for path in export_paths:
                    if os.path.isdir(path):
                        # 导出文件夹
                        self.export_folder_to_pdf(c, path, width, height)
                    else:
                        # 导出单个文件
                        self.export_file_to_pdf(c, path, width, height)
                
                c.save()
                self.root.after(0, lambda: messagebox.showinfo("导出完成", f"PDF已导出到: {output_file}"))
                self.root.after(0, lambda: self.status_var.set("PDF导出完成"))
                
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("导出错误", f"PDF导出失败: {str(e)}"))
                self.root.after(0, lambda: self.status_var.set("PDF导出失败"))
            finally:
                self.root.after(0, lambda: self.progress.stop())
        
        thread = threading.Thread(target=export_thread)
        thread.daemon = True
        thread.start()
    
    def export_folder_to_pdf(self, canvas_obj, folder_path, page_width, page_height):
        """导出文件夹到PDF"""
        items = os.listdir(folder_path)
        items.sort()
        
        for item in items:
            item_path = os.path.join(folder_path, item)
            
            if os.path.isdir(item_path):
                # 子文件夹
                self.export_folder_to_pdf(canvas_obj, item_path, page_width, page_height)
            else:
                # 文件
                self.export_file_to_pdf(canvas_obj, item_path, page_width, page_height)
    
    def export_file_to_pdf(self, canvas_obj, file_path, page_width, page_height):
        """导出单个文件到PDF"""
        filename = os.path.basename(file_path)
        
        if file_path.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif')):
            # 图片文件
            try:
                # 添加标题
                canvas_obj.setFont("Helvetica-Bold", 16)
                canvas_obj.drawString(50, page_height - 50, filename)
                
                # 添加图片
                img = ImageReader(file_path)
                img_width, img_height = img.getSize()
                
                # 计算图片显示尺寸
                max_width = page_width - 100
                max_height = page_height - 150
                
                ratio = min(max_width/img_width, max_height/img_height)
                display_width = img_width * ratio
                display_height = img_height * ratio
                
                x = (page_width - display_width) / 2
                y = (page_height - display_height) / 2 - 50
                
                canvas_obj.drawImage(img, x, y, width=display_width, height=display_height)
                
                # 检查是否有OCR文本
                base_name = os.path.splitext(filename)[0]
                ocr_file = os.path.join(os.path.dirname(file_path), f"{base_name}_ocr.txt")
                
                if os.path.exists(ocr_file):
                    # 添加OCR文本
                    try:
                        with open(ocr_file, 'r', encoding='utf-8') as f:
                            ocr_text = f.read()
                        
                        canvas_obj.showPage()
                        canvas_obj.setFont("Helvetica", 12)
                        canvas_obj.drawString(50, page_height - 50, f"OCR识别结果 - {filename}")
                        
                        # 文本换行处理
                        lines = ocr_text.split('\n')
                        y_position = page_height - 80
                        
                        for line in lines:
                            if y_position < 50:
                                canvas_obj.showPage()
                                canvas_obj.setFont("Helvetica", 12)
                                y_position = page_height - 50
                            
                            canvas_obj.drawString(50, y_position, line[:100])  # 限制每行长度
                            y_position -= 20
                    
                    except Exception as e:
                        print(f"读取OCR文件失败: {e}")
                
                canvas_obj.showPage()
                
            except Exception as e:
                print(f"导出图片失败: {e}")
        
        elif file_path.lower().endswith('.txt'):
            # 文本文件
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                canvas_obj.setFont("Helvetica-Bold", 16)
                canvas_obj.drawString(50, page_height - 50, filename)
                
                canvas_obj.setFont("Helvetica", 12)
                lines = content.split('\n')
                y_position = page_height - 80
                
                for line in lines:
                    if y_position < 50:
                        canvas_obj.showPage()
                        canvas_obj.setFont("Helvetica", 12)
                        y_position = page_height - 50
                    
                    canvas_obj.drawString(50, y_position, line[:100])
                    y_position -= 20
                
                canvas_obj.showPage()
                
            except Exception as e:
                print(f"导出文本文件失败: {e}")
    
    def export_word(self):
        """导出Word文档"""
        # 选择导出范围（与PDF类似）
        export_window = tk.Toplevel(self.root)
        export_window.title("导出Word")
        export_window.geometry("400x300")
        export_window.transient(self.root)
        export_window.grab_set()
        
        ttk.Label(export_window, text="选择导出范围:").pack(pady=10)
        
        export_var = tk.StringVar(value="current")
        
        ttk.Radiobutton(export_window, text="当前目录所有文件", 
                       variable=export_var, value="current").pack(anchor=tk.W, padx=20, pady=5)
        ttk.Radiobutton(export_window, text="按学科导出", 
                       variable=export_var, value="subject").pack(anchor=tk.W, padx=20, pady=5)
        ttk.Radiobutton(export_window, text="选择文件和文件夹", 
                       variable=export_var, value="selected").pack(anchor=tk.W, padx=20, pady=5)
        
        subject_frame = ttk.Frame(export_window)
        subject_frame.pack(fill=tk.X, padx=20, pady=5)
        
        ttk.Label(subject_frame, text="学科:").pack(side=tk.LEFT)
        subject_var = tk.StringVar()
        subject_combo = ttk.Combobox(subject_frame, textvariable=subject_var, 
                                   values=self.subjects, width=20)
        subject_combo.pack(side=tk.LEFT, padx=5)
        
        def on_export_option_change(*args):
            if export_var.get() == "subject":
                subject_frame.pack(fill=tk.X, padx=20, pady=5)
            else:
                subject_frame.pack_forget()
        
        export_var.trace('w', on_export_option_change)
        
        button_frame = ttk.Frame(export_window)
        button_frame.pack(fill=tk.X, padx=20, pady=20)
        
        def do_export():
            export_type = export_var.get()
            
            if export_type == "current":
                export_paths = [self.current_path]
            elif export_type == "subject":
                subject = subject_var.get()
                if not subject:
                    messagebox.showwarning("警告", "请选择学科")
                    return
                export_paths = [os.path.join(self.cuoti_dir, subject)]
            else:
                export_paths = filedialog.askopenfilenames(
                    title="选择要导出的文件",
                    initialdir=self.current_path,
                    parent=export_window
                )
                if not export_paths:
                    return
            
            output_file = filedialog.asksaveasfilename(
                title="保存Word文档",
                defaultextension=".docx",
                filetypes=[("Word文档", "*.docx")],
                parent=export_window
            )
            
            if not output_file:
                return
            
            export_window.destroy()
            self.perform_word_export(export_paths, output_file)
        
        ttk.Button(button_frame, text="确定", command=do_export).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=export_window.destroy).pack(side=tk.RIGHT, padx=5)
    
    def perform_word_export(self, export_paths, output_file):
        """执行Word导出"""
        self.status_var.set("正在导出Word...")
        self.progress.start()
        
        def export_thread():
            try:
                doc = Document()
                
                for path in export_paths:
                    if os.path.isdir(path):
                        # 导出文件夹
                        self.export_folder_to_word(doc, path)
                    else:
                        # 导出单个文件
                        self.export_file_to_word(doc, path)
                
                doc.save(output_file)
                self.root.after(0, lambda: messagebox.showinfo("导出完成", f"Word文档已导出到: {output_file}"))
                self.root.after(0, lambda: self.status_var.set("Word导出完成"))
                
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("导出错误", f"Word导出失败: {str(e)}"))
                self.root.after(0, lambda: self.status_var.set("Word导出失败"))
            finally:
                self.root.after(0, lambda: self.progress.stop())
        
        thread = threading.Thread(target=export_thread)
        thread.daemon = True
        thread.start()
    
    def export_folder_to_word(self, doc, folder_path):
        """导出文件夹到Word"""
        items = os.listdir(folder_path)
        items.sort()
        
        folder_name = os.path.basename(folder_path)
        doc.add_heading(f'文件夹: {folder_name}', level=1)
        
        for item in items:
            item_path = os.path.join(folder_path, item)
            
            if os.path.isdir(item_path):
                # 子文件夹
                self.export_folder_to_word(doc, item_path)
            else:
                # 文件
                self.export_file_to_word(doc, item_path)
    
    def export_file_to_word(self, doc, file_path):
        """导出单个文件到Word"""
        filename = os.path.basename(file_path)
        
        if file_path.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif')):
            # 图片文件
            try:
                doc.add_heading(filename, level=2)
                
                # 添加图片
                doc.add_picture(file_path, width=Inches(6))
                
                # 检查是否有OCR文本
                base_name = os.path.splitext(filename)[0]
                ocr_file = os.path.join(os.path.dirname(file_path), f"{base_name}_ocr.txt")
                
                if os.path.exists(ocr_file):
                    # 添加OCR文本
                    try:
                        with open(ocr_file, 'r', encoding='utf-8') as f:
                            ocr_text = f.read()
                        
                        doc.add_heading('OCR识别结果', level=3)
                        doc.add_paragraph(ocr_text)
                    except Exception as e:
                        print(f"读取OCR文件失败: {e}")
                
                doc.add_page_break()
                
            except Exception as e:
                print(f"导出图片失败: {e}")
        
        elif file_path.lower().endswith('.txt'):
            # 文本文件
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                doc.add_heading(filename, level=2)
                doc.add_paragraph(content)
                doc.add_page_break()
                
            except Exception as e:
                print(f"导出文本文件失败: {e}")
    
    def export_selected(self):
        """导出选中的项目"""
        selection = self.file_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择要导出的项目")
            return
        
        export_format = simpledialog.askstring("导出格式", "请输入导出格式 (pdf/word):", 
                                             initialvalue="pdf", parent=self.root)
        if not export_format or export_format.lower() not in ['pdf', 'word']:
            return
        
        export_paths = []
        for item in selection:
            filename = self.file_tree.item(item, "values")[0]
            if filename != "返回上一级":
                export_paths.append(os.path.join(self.current_path, filename))
        
        if not export_paths:
            return
        
        if export_format.lower() == 'pdf':
            output_file = filedialog.asksaveasfilename(
                title="保存PDF文件",
                defaultextension=".pdf",
                filetypes=[("PDF文件", "*.pdf")],
                parent=self.root
            )
            if output_file:
                self.perform_pdf_export(export_paths, output_file)
        else:
            output_file = filedialog.asksaveasfilename(
                title="保存Word文档",
                defaultextension=".docx",
                filetypes=[("Word文档", "*.docx")],
                parent=self.root
            )
            if output_file:
                self.perform_word_export(export_paths, output_file)
    
    def batch_process(self):
        """批量处理"""
        messagebox.showinfo("功能开发中", "批量处理功能正在开发中，敬请期待！")
    
    def show_settings(self):
        """显示设置窗口"""
        settings_window = tk.Toplevel(self.root)
        settings_window.title("设置")
        settings_window.geometry("400x300")
        settings_window.transient(self.root)
        settings_window.grab_set()
        
        # OCR设置
        ocr_frame = ttk.LabelFrame(settings_window, text="OCR设置")
        ocr_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ocr_enabled_var = tk.BooleanVar(value=self.config.get("ocr_enabled", True))
        ttk.Checkbutton(ocr_frame, text="启用OCR功能", 
                       variable=ocr_enabled_var).pack(anchor=tk.W, padx=5, pady=2)
        
        # 图片质量设置
        quality_frame = ttk.LabelFrame(settings_window, text="图片质量")
        quality_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(quality_frame, text="保存质量 (1-100):").pack(anchor=tk.W, padx=5)
        quality_var = tk.IntVar(value=self.config.get("image_quality", 90))
        quality_scale = ttk.Scale(quality_frame, from_=10, to=100, 
                                variable=quality_var, orient=tk.HORIZONTAL)
        quality_scale.pack(fill=tk.X, padx=5, pady=2)
        
        quality_label = ttk.Label(quality_frame, text=f"当前质量: {quality_var.get()}")
        quality_label.pack(anchor=tk.W, padx=5)
        
        def on_quality_change(*args):
            quality_label.config(text=f"当前质量: {quality_var.get()}")
        
        quality_var.trace('w', on_quality_change)
        
        # 按钮
        button_frame = ttk.Frame(settings_window)
        button_frame.pack(fill=tk.X, padx=10, pady=20)
        
        def save_settings():
            self.config["ocr_enabled"] = ocr_enabled_var.get()
            self.config["image_quality"] = quality_var.get()
            self.save_config()
            messagebox.showinfo("保存成功", "设置已保存")
            settings_window.destroy()
        
        ttk.Button(button_frame, text="保存", command=save_settings).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=settings_window.destroy).pack(side=tk.RIGHT, padx=5)
    
    def show_help(self):
        """显示帮助信息"""
        help_text = """
错题整理工具使用说明

主要功能：
1. 导入错题：点击"导入错题"按钮，选择学科和图片文件
2. 文件管理：双击文件夹进入，双击文件预览，右键菜单进行操作
3. 图片处理：选择图片后点击"图片处理"进行调整
4. OCR识别：选择图片后点击"OCR识别"提取文字
5. 导出功能：支持导出为PDF和Word格式

操作技巧：
- 双击文件夹进入下一级目录
- 双击文件进行预览
- 右键点击项目显示操作菜单
- 使用"返回上一级"按钮返回上级目录
- 在编辑区域可以修改题目信息和添加备注

快捷键：
- Ctrl+I: 导入错题
- Ctrl+R: 刷新列表
- Ctrl+E: 导出PDF
- F5: 刷新列表

作者：mmm
版本：1.0.0
赞助链接：https://gitee.com/orangearc655743/Wrong-Question-Tool.git
        """
        
        help_window = tk.Toplevel(self.root)
        help_window.title("使用说明")
        help_window.geometry("600x500")
        help_window.transient(self.root)
        
        # 文本显示
        text_widget = tk.Text(help_window, wrap=tk.WORD, padx=10, pady=10)
        text_widget.pack(fill=tk.BOTH, expand=True)
        text_widget.insert(1.0, help_text)
        text_widget.config(state=tk.DISABLED)
        
        # 滚动条
        scrollbar = ttk.Scrollbar(help_window, orient=tk.VERTICAL, command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        ttk.Button(help_window, text="关闭", 
                  command=help_window.destroy).pack(pady=10)
    
    def show_about(self):
        """显示关于信息"""
        about_text = """
错题整理工具 v2.0.0

一个功能强大且美观的错题整理和管理工具

主要功能：
✓ 导入和管理错题图片
✓ OCR文字识别
✓ 图片预处理、增强、裁剪和旋转
✓ PDF和Word导出
✓ 学科分类管理
✓ 标签系统
✓ 批量重命名
✓ 文件搜索
✓ 统计信息
✓ 快捷键支持
✓ 主题切换
✓ 跨平台支持

新增功能：
🔧 图片裁剪功能
🔧 图片旋转功能
🔧 标签管理系统
🔧 批量重命名
🔧 文件搜索
🔧 统计信息显示
🔧 快捷键支持
🔧 界面美化
🔧 自动备份

技术栈：
- Python 3.x
- Tkinter (GUI)
- PIL (图片处理)
- OpenCV (图像处理)
- pytesseract (OCR)
- reportlab (PDF生成)
- python-docx (Word生成)

作者：mmm
版本：2.0.0
发布日期：2025-11-02

赞助链接：https://gitee.com/orangearc655743/Wrong-Question-Tool.git

感谢您的使用和支持！
        """
        
        messagebox.showinfo("关于", about_text)
    
    def open_sponsor_link(self):
        """打开赞助链接"""
        webbrowser.open("https://gitee.com/orangearc655743/Wrong-Question-Tool.git")
    
    def add_tags(self):
        """添加标签"""
        selection = self.file_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一个文件")
            return
        
        item = selection[0]
        filename = self.file_tree.item(item, "values")[0]
        
        if filename == "返回上一级":
            return
        
        # 获取当前标签
        current_tags = self.get_file_tags(filename)
        
        # 标签输入窗口
        tag_window = tk.Toplevel(self.root)
        tag_window.title(f"添加标签 - {filename}")
        tag_window.geometry("400x200")
        tag_window.transient(self.root)
        tag_window.grab_set()
        
        ttk.Label(tag_window, text="标签 (用逗号分隔):", style='Header.TLabel').pack(pady=10)
        
        tag_text = tk.Text(tag_window, height=8, width=40)
        tag_text.pack(pady=10, padx=20, fill=tk.BOTH, expand=True)
        tag_text.insert(1.0, current_tags)
        
        def save_tags():
            try:
                new_tags = tag_text.get(1.0, tk.END).strip()
                
                # 保存标签到元数据文件
                base_name = os.path.splitext(filename)[0]
                meta_file = os.path.join(self.current_path, f"{base_name}.meta")
                
                metadata = {"tags": new_tags}
                
                if os.path.exists(meta_file):
                    # 读取现有元数据
                    with open(meta_file, 'r', encoding='utf-8') as f:
                        existing_metadata = json.load(f)
                    metadata.update(existing_metadata)
                
                with open(meta_file, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, ensure_ascii=False, indent=2)
                
                messagebox.showinfo("成功", "标签已保存")
                tag_window.destroy()
                self.refresh_file_list()
                
            except Exception as e:
                messagebox.showerror("错误", f"保存标签失败: {str(e)}")
        
        button_frame = ttk.Frame(tag_window)
        button_frame.pack(pady=10)
        
        ttk.Button(button_frame, text="保存", command=save_tags).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=tag_window.destroy).pack(side=tk.LEFT, padx=5)
    
    def edit_tags(self):
        """编辑标签（与添加标签相同功能）"""
        self.add_tags()
    
    def batch_rename(self):
        """批量重命名"""
        selection = self.file_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择要重命名的文件")
            return
        
        # 重命名窗口
        rename_window = tk.Toplevel(self.root)
        rename_window.title("批量重命名")
        rename_window.geometry("500x400")
        rename_window.transient(self.root)
        rename_window.grab_set()
        
        ttk.Label(rename_window, text="批量重命名设置:", style='Header.TLabel').pack(pady=10)
        
        # 重命名选项
        options_frame = ttk.Frame(rename_window)
        options_frame.pack(fill=tk.X, padx=20, pady=10)
        
        ttk.Label(options_frame, text="前缀:").grid(row=0, column=0, sticky=tk.W, pady=5)
        prefix_var = tk.StringVar()
        ttk.Entry(options_frame, textvariable=prefix_var, width=20).grid(row=0, column=1, sticky=tk.W+tk.E, pady=5)
        
        ttk.Label(options_frame, text="后缀:").grid(row=1, column=0, sticky=tk.W, pady=5)
        suffix_var = tk.StringVar()
        ttk.Entry(options_frame, textvariable=suffix_var, width=20).grid(row=1, column=1, sticky=tk.W+tk.E, pady=5)
        
        ttk.Label(options_frame, text="起始编号:").grid(row=2, column=0, sticky=tk.W, pady=5)
        start_num_var = tk.IntVar(value=1)
        ttk.Spinbox(options_frame, from_=1, to=9999, textvariable=start_num_var, width=18).grid(row=2, column=1, sticky=tk.W+tk.E, pady=5)
        
        ttk.Label(options_frame, text="编号位数:").grid(row=3, column=0, sticky=tk.W, pady=5)
        digit_count_var = tk.IntVar(value=3)
        ttk.Spinbox(options_frame, from_=1, to=10, textvariable=digit_count_var, width=18).grid(row=3, column=1, sticky=tk.W+tk.E, pady=5)
        
        # 预览列表
        preview_frame = ttk.LabelFrame(rename_window, text="预览")
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        preview_text = tk.Text(preview_frame, height=10)
        preview_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        def update_preview():
            preview_text.delete(1.0, tk.END)
            counter = start_num_var.get()
            
            for item in selection:
                filename = self.file_tree.item(item, "values")[0]
                if filename != "返回上一级":
                    name, ext = os.path.splitext(filename)
                    new_name = f"{prefix_var.get()}{name}{suffix_var.get()}_{counter:0{digit_count_var.get()}d}{ext}"
                    preview_text.insert(tk.END, f"{filename} -> {new_name}\n")
                    counter += 1
        
        # 绑定预览更新
        prefix_var.trace('w', lambda *args: update_preview())
        suffix_var.trace('w', lambda *args: update_preview())
        start_num_var.trace('w', lambda *args: update_preview())
        digit_count_var.trace('w', lambda *args: update_preview())
        
        # 按钮
        button_frame = ttk.Frame(rename_window)
        button_frame.pack(pady=10)
        
        def apply_rename():
            try:
                counter = start_num_var.get()
                renamed_count = 0
                
                for item in selection:
                    filename = self.file_tree.item(item, "values")[0]
                    if filename != "返回上一级":
                        old_path = os.path.join(self.current_path, filename)
                        name, ext = os.path.splitext(filename)
                        new_name = f"{prefix_var.get()}{name}{suffix_var.get()}_{counter:0{digit_count_var.get()}d}{ext}"
                        new_path = os.path.join(self.current_path, new_name)
                        
                        if not os.path.exists(new_path):
                            os.rename(old_path, new_path)
                            renamed_count += 1
                        
                        counter += 1
                
                messagebox.showinfo("完成", f"成功重命名 {renamed_count} 个文件")
                rename_window.destroy()
                self.refresh_file_list()
                
            except Exception as e:
                messagebox.showerror("错误", f"重命名失败: {str(e)}")
        
        ttk.Button(button_frame, text="应用重命名", command=apply_rename).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=rename_window.destroy).pack(side=tk.LEFT, padx=5)
        
        # 配置grid权重
        options_frame.columnconfigure(1, weight=1)
        
        # 初始化预览
        update_preview()
    
    def create_backup(self):
        """创建备份"""
        try:
            backup_dir = os.path.join(self.program_dir, "backup", datetime.datetime.now().strftime("%Y%m%d_%H%M%S"))
            os.makedirs(backup_dir, exist_ok=True)
            
            # 复制所有文件
            if os.path.exists(self.cuoti_dir):
                for item in os.listdir(self.cuoti_dir):
                    src = os.path.join(self.cuoti_dir, item)
                    dst = os.path.join(backup_dir, item)
                    if os.path.isdir(src):
                        shutil.copytree(src, dst)
                    else:
                        shutil.copy2(src, dst)
            
            messagebox.showinfo("备份完成", f"备份已创建: {backup_dir}")
            
        except Exception as e:
            messagebox.showerror("错误", f"创建备份失败: {str(e)}")
    
    def show_statistics(self):
        """显示统计信息"""
        stats_window = tk.Toplevel(self.root)
        stats_window.title("统计信息")
        stats_window.geometry("600x500")
        stats_window.transient(self.root)
        
        # 总体统计
        total_frame = ttk.LabelFrame(stats_window, text="总体统计")
        total_frame.pack(fill=tk.X, padx=10, pady=5)
        
        total_text = f"总文件数: {self.stats['total_files']}\n"
        
        total_size = self.stats['total_size']
        if total_size < 1024:
            size_str = f"{total_size} B"
        elif total_size < 1024*1024:
            size_str = f"{total_size/1024:.1f} KB"
        elif total_size < 1024*1024*1024:
            size_str = f"{total_size/(1024*1024):.1f} MB"
        else:
            size_str = f"{total_size/(1024*1024*1024):.1f} GB"
        
        total_text += f"总大小: {size_str}\n"
        total_text += f"学科数量: {len(self.stats['by_subject'])}"
        
        ttk.Label(total_frame, text=total_text, style='Header.TLabel').pack(pady=10)
        
        # 按学科统计
        subject_frame = ttk.LabelFrame(stats_window, text="按学科统计")
        subject_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 创建Treeview显示学科统计
        subject_tree = ttk.Treeview(subject_frame, columns=("数量",), show="tree headings")
        subject_tree.heading("#0", text="学科")
        subject_tree.heading("数量", text="文件数量")
        subject_tree.column("#0", width=200)
        subject_tree.column("数量", width=100)
        
        for subject, count in sorted(self.stats['by_subject'].items()):
            subject_tree.insert("", "end", text=subject, values=(count,))
        
        subject_tree.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        ttk.Button(stats_window, text="关闭", command=stats_window.destroy).pack(pady=10)
    
    def show_search(self):
        """显示搜索窗口"""
        self.focus_search()
    
    def change_view_mode(self, mode):
        """更改视图模式"""
        # 这里可以实现不同的视图模式，如大图标、小图标等
        messagebox.showinfo("功能提示", f"视图模式 '{mode}' 正在开发中")
    
    def switch_theme(self):
        """切换主题"""
        themes = ["默认", "深色", "浅色"]
        theme = simpledialog.askstring("选择主题", f"可用主题:\n" + "\n".join(themes), 
                                     initialvalue="默认", parent=self.root)
        if theme:
            self.config["theme"] = theme
            self.save_config()
            messagebox.showinfo("主题切换", f"主题已切换为: {theme}\n重启程序后生效")
    
    def show_shortcuts(self):
        """显示快捷键"""
        shortcuts_text = """
快捷键列表

文件操作:
Ctrl+I - 导入错题
Ctrl+R - 刷新列表
Ctrl+E - 导出PDF
F5 - 刷新列表

编辑操作:
F2 - 重命名
Delete - 删除

搜索:
Ctrl+F - 搜索文件

通用:
ESC - 取消当前操作
        """
        
        shortcuts_window = tk.Toplevel(self.root)
        shortcuts_window.title("快捷键")
        shortcuts_window.geometry("400x300")
        shortcuts_window.transient(self.root)
        
        text_widget = tk.Text(shortcuts_window, wrap=tk.WORD, padx=10, pady=10)
        text_widget.pack(fill=tk.BOTH, expand=True)
        text_widget.insert(1.0, shortcuts_text)
        text_widget.config(state=tk.DISABLED)
        
        ttk.Button(shortcuts_window, text="关闭", 
                  command=shortcuts_window.destroy).pack(pady=10)
    
    def run(self):
        """运行程序"""
        try:
            self.root.mainloop()
        except KeyboardInterrupt:
            print("\n程序被用户中断")
        except Exception as e:
            print(f"程序运行错误: {e}")
            messagebox.showerror("程序错误", f"程序运行出现错误: {str(e)}")

def main():
    """主函数"""
    print("错题整理工具 v2.0.0")
    print("作者：mmm")
    print("赞助链接：https://gitee.com/orangearc655743/Wrong-Question-Tool.git")
    print("正在启动...")
    
    try:
        app = WrongQuestionTool()
        app.run()
    except Exception as e:
        print(f"程序启动失败: {e}")
        input("按回车键退出...")

if __name__ == "__main__":
    main()